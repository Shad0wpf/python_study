#coding=utf-8

from docx import Document

doc = Document('test.docx')

for i in doc.paragraphs:
    print i.text

doc.add_paragraph('555')

# doc.add_heading('Title111', level=2)
# doc.add_heading('Title222', level=3)

doc.save('test2.docx')